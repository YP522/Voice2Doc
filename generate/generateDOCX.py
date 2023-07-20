import docx
import os
# from docx import Document
from system import utils as u

def makeDoc(name,txt):
    doc = docx.Document()
    doc.add_paragraph(txt)
    doc.save(u.dt_string+"/docs_chunks/"f'{name}.docx')

def get_docx_files(path):
    files = []
    for file in os.listdir(path):
        if file.endswith(".docx"):
            files.append(os.path.join(path, file))
    return files

def get_docx_text(file):
    doc = docx.Document(file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)



def combine_word_documents(files):
    merged_document = docx.Document()

    for index, file in enumerate(files):
        sub_doc = docx.Document(file)

        # Don't add a page break if you've reached the last file.
        if index < len(files)-1:
           sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    merged_document.save(u.dt_string+'/merged.docx')

  